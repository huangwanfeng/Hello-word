# -*- coding: utf-8 -*-
from docx import Document
from PIL import Image, ImageDraw
from io import BytesIO


document = Document()#新建一个文档
p = document.add_paragraph()#添加一个段落
r = p.add_run()#添加一个游程？？什么是游程
img_size = 20

for x in range(255):
    im = Image.new("RGB", (img_size,img_size),"white") #图形颜色模式、尺寸、填充色
    draw_obj = ImageDraw.Draw(im)
    draw_obj.ellipse((0,0,img_size-1,img_size-1), fill = 255 - x, outline = x) #画椭圆，坐标，尺寸、填充色、线条色
    fake_buf_file = BytesIO()# 用BytesIO将图片保存在内存里，减少磁盘操作
    im.save(fake_buf_file,'png') #图片格式？
    r.add_picture(fake_buf_file)#在当前游程中插入图片
    fake_buf_file.close()
    document.save("demo.docx")